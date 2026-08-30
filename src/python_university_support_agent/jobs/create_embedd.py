
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from ..config import settings
from langchain_text_splitters import RecursiveCharacterTextSplitter
import pypdf
from docx import Document as DocxDocument
from langchain_core.documents import Document
from ..retrieval import get_vector_store, get_embeddings
from pathlib import Path
from io import BytesIO
from ..db import sessionLocal
from ..models import Document as DocumentModel
from ..logger import get_logger
from tenacity import retry, retry_if_exception_type, wait_exponential, stop_after_attempt
import asyncio
from python_university_support_agent.services import start_job, complete_job, update_progress, fail_job


BATCH_SIZE = 10

logger = get_logger("Embedd Job")

text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,      # max characters per chunk
    chunk_overlap=200,    # overlap between chunks (preserves context across boundaries)
    add_start_index=True  # track index of chunk in original document
)

embeddings = get_embeddings()

vector_store = get_vector_store( embeddings )



async def get_documents_from_file(file_abs_path: Path | str, filename: str) -> list[Document]:
    file_path = file_abs_path if isinstance(file_abs_path, Path) else Path(file_abs_path)
    ext = file_path.suffix
    docs = list()

    if ext == ".pdf":
        reader = pypdf.PdfReader(file_path)
        for i, page in enumerate(reader.pages):
            page_content = page.extract_text()
            docs.append(Document(page_content=page_content or "", metadata = { "source": filename, "page": i}))
    elif ext == ".docx":
        docx = DocxDocument( BytesIO( file_path.read_bytes() ))
        for index, paragraph in enumerate(docx.paragraphs, start=1):
            docs.append(Document(page_content=paragraph.text.strip() or "", metadata = { "source": filename, "paragraph": index}))
    elif ext == ".txt":
        content = file_path.read_text(encoding="utf8")
        docs.append(Document(page_content= content, metadata= {"source": filename, "page": 1}))
    return docs


@retry(
       retry=retry_if_exception_type((TimeoutError, asyncio.TimeoutError, ConnectionError)),
       stop=stop_after_attempt(3),
       wait=wait_exponential(
            multiplier=1,
            min=2,
            max=60,
      ),
      reraise=True,
)
async def process_batch(batch: list[Document]):
    await asyncio.wait_for(
        vector_store.aadd_documents(batch),
        timeout= 60 * 5
    )

async def create_embedd(ctx, job_id):
    from ..logger import get_logger

    logger = get_logger(f"Job<{job_id}> | Embedd Task")
    async with sessionLocal() as db:
        try:
            job = await start_job(db=db, job_id=job_id)
            doc = job.document

            logger = get_logger(f"Doc<{doc.filename}> | Embedd Task")
            logger.info("Task started")

            filename = doc.filename
            file_abs_path = Path.joinpath(settings.storage_dir, filename)

            docs: list[Document] = await get_documents_from_file(file_abs_path=file_abs_path, filename=filename)
            logger.info(f"Total {len(docs)} langchain docs will be processed: ")

            all_splits = text_splitter.split_documents(docs)
            total_docs = len(all_splits)

            await update_progress(db=db, job_id=job.id, embedded_chunks=0, total_chunks=total_docs)

            for i in range(0, total_docs, BATCH_SIZE):
                start = i
                end = min(start + BATCH_SIZE, total_docs)
                batch = all_splits[start:end]
                logger.info(
                    "Batch processing started | range=%d-%d | total_documents=%d",
                    start,
                    end - 1,
                    total_docs
                )
                await process_batch(batch=batch)
                await update_progress(db=db, job_id=job.id, embedded_chunks=end)
                logger.info(
                    "Batch processing ended | range=%d-%d | total_documents=%d",
                    start,
                    end - 1,
                    total_docs
                )

            await complete_job(db=db, job_id=job.id)
            logger.info(f"File {filename} embedding has processed successfully")
        except Exception as e:
            logger.error(f"Embedding job failed: {e}")
            await fail_job(db=db, job_id=job_id, error=str(e))
            raise e




